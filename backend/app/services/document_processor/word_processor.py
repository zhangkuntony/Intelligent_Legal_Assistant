"""
Word文档处理器
"""

from typing import Dict, Any, List, Optional
from dataclasses import dataclass

from .base_processor import BaseDocumentProcessor, DocumentChunk
from .exceptions import ExtractionError

try:
    from docx import Document as DocxDocument

    WORD_SUPPORT_AVAILABLE = True
except ImportError:
    WORD_SUPPORT_AVAILABLE = False


@dataclass
class StyleInfo:
    """样式信息"""
    style_name: str
    font_name: Optional[str]
    font_size: Optional[float]
    is_bold: bool
    is_italic: bool
    is_underline: bool
    alignment: Optional[str]


@dataclass
class TableInfo:
    """表格信息"""
    table_index: int
    rows: int
    columns: int
    content: str


class WordProcessor(BaseDocumentProcessor):
    """Word文档处理器"""

    def __init__(self, config: Dict[str, Any] = None):
        super().__init__(config)
        self.supported_formats = ['.docx', '.doc']

        if not WORD_SUPPORT_AVAILABLE:
            raise ImportError("Word处理依赖未安装，请安装 python-docx")

    async def extract_text(self, file_path: str) -> str:
        """提取Word文档文本内容"""
        try:
            doc = DocxDocument(file_path)
            full_text = []

            # 提取段落文本
            for para in doc.paragraphs:
                if para.text.strip():
                    full_text.append(para.text)

            # 提取表格文本
            for table_index, table in enumerate(doc.tables):
                table_text = self._extract_table_text(table, table_index)
                if table_text:
                    full_text.append(f"\n--- 表格{table_index + 1} ---\n")
                    full_text.append(table_text)

            return "\n".join(full_text)

        except Exception as e:
            raise ExtractionError(f"Word文档文本提取失败: {str(e)}")

    async def extract_metadata(self, file_path: str) -> Dict[str, Any]:
        """提取Word文档元数据"""
        try:
            metadata = await super().extract_metadata(file_path)

            doc = DocxDocument(file_path)
            core_props = doc.core_properties

            # 文档属性
            word_metadata = {
                'title': core_props.title,
                'author': core_props.author,
                'subject': core_props.subject,
                'keywords': core_props.keywords,
                'comments': core_props.comments,
                'last_modified_by': core_props.last_modified_by,
                'revision': core_props.revision,
                'created': core_props.created,
                'modified': core_props.modified,
                'last_printed': core_props.last_printed,
                'category': core_props.category,
                'content_status': core_props.content_status,
            }

            # 清理空值
            word_metadata = {k: v for k, v in word_metadata.items() if v}

            metadata.update(word_metadata)
            return metadata

        except Exception as e:
            # 元数据提取失败不影响主要功能
            return await super().extract_metadata(file_path)

    async def extract_styles(self, file_path: str) -> List[StyleInfo]:
        """提取样式信息"""
        styles = []

        try:
            doc = DocxDocument(file_path)

            # 分析文档中的样式使用情况
            style_usage = {}

            for para in doc.paragraphs:
                if para.style and para.text.strip():
                    style_name = para.style.name

                    if style_name not in style_usage:
                        # 创建样式信息
                        style_info = StyleInfo(
                            style_name=style_name,
                            font_name=None,
                            font_size=None,
                            is_bold=False,
                            is_italic=False,
                            is_underline=False,
                            alignment=None
                        )

                        # 尝试获取字体信息
                        if para.runs:
                            first_run = para.runs[0]
                            style_info.font_name = first_run.font.name
                            style_info.font_size = first_run.font.size
                            style_info.is_bold = first_run.font.bold
                            style_info.is_italic = first_run.font.italic
                            style_info.is_underline = first_run.font.underline

                        # 获取对齐方式
                        if para.alignment:
                            style_info.alignment = str(para.alignment)

                        styles.append(style_info)
                        style_usage[style_name] = style_info

        except Exception as e:
            # 样式提取失败不影响主要功能
            pass

        return styles

    async def extract_tables(self, file_path: str) -> List[TableInfo]:
        """提取表格数据"""
        tables = []

        try:
            doc = DocxDocument(file_path)

            for table_index, table in enumerate(doc.tables):
                table_content = self._extract_table_content(table)

                table_info = TableInfo(
                    table_index=table_index,
                    rows=len(table.rows),
                    columns=len(table.columns) if table.columns else 0,
                    content=table_content
                )

                tables.append(table_info)

        except Exception as e:
            # 表格提取失败不影响主要功能
            pass

        return tables

    async def chunk_text(self, text: str, **kwargs) -> List[DocumentChunk]:
        """Word文档专用分块策略"""
        chunks = []
        chunk_index = 0

        # 按段落分割
        paragraphs = text.split('\n')
        current_chunk = ""
        current_heading = ""

        max_chunk_size = kwargs.get('chunk_size', self.config.get('default_chunk_size', 1000))

        for paragraph in paragraphs:
            paragraph = paragraph.strip()
            if not paragraph:
                continue

            # 检查是否是标题或分隔符
            if paragraph.startswith('--- 表格') and paragraph.endswith('---'):
                # 表格分隔符，强制分割
                if current_chunk:
                    chunk = DocumentChunk(
                        chunk_id=f"word_chunk_{chunk_index}",
                        content=current_chunk,
                        chunk_index=chunk_index,
                        metadata={
                            'chunk_type': 'paragraph_group',
                            'heading': current_heading,
                            'content_length': len(current_chunk)
                        },
                        section_title=current_heading
                    )
                    chunks.append(chunk)
                    chunk_index += 1
                    current_chunk = ""

                # 表格作为独立块
                continue

            # 检查是否是标题（简单的启发式规则）
            is_heading = self._is_heading(paragraph)

            if is_heading:
                current_heading = paragraph

                # 如果当前块不为空，先保存当前块
                if current_chunk:
                    chunk = DocumentChunk(
                        chunk_id=f"word_chunk_{chunk_index}",
                        content=current_chunk,
                        chunk_index=chunk_index,
                        metadata={
                            'chunk_type': 'paragraph_group',
                            'heading': current_heading,
                            'content_length': len(current_chunk)
                        },
                        section_title=current_heading
                    )
                    chunks.append(chunk)
                    chunk_index += 1
                    current_chunk = ""

            # 添加段落到当前块
            if current_chunk:
                current_chunk += "\n" + paragraph
            else:
                current_chunk = paragraph

            # 检查是否超过分块大小限制
            if len(current_chunk) > max_chunk_size:
                # 强制分割当前块
                chunk_content = current_chunk[:max_chunk_size]
                remaining_content = current_chunk[max_chunk_size:]

                chunk = DocumentChunk(
                    chunk_id=f"word_chunk_{chunk_index}",
                    content=chunk_content,
                    chunk_index=chunk_index,
                    metadata={
                        'chunk_type': 'paragraph_group',
                        'heading': current_heading,
                        'content_length': len(chunk_content),
                        'is_truncated': True
                    },
                    section_title=current_heading
                )
                chunks.append(chunk)
                chunk_index += 1

                # 剩余内容作为新块开始
                current_chunk = remaining_content

        # 添加最后一个块
        if current_chunk:
            chunk = DocumentChunk(
                chunk_id=f"word_chunk_{chunk_index}",
                content=current_chunk,
                chunk_index=chunk_index,
                metadata={
                    'chunk_type': 'paragraph_group',
                    'heading': current_heading,
                    'content_length': len(current_chunk)
                },
                section_title=current_heading
            )
            chunks.append(chunk)

        return chunks

    def _extract_table_text(self, table, table_index: int) -> str:
        """提取表格文本"""
        table_content = []

        for row in table.rows:
            row_content = []
            for cell in row.cells:
                cell_text = cell.text.strip()
                row_content.append(cell_text)

            if any(row_content):  # 只添加非空行
                table_content.append(" | ".join(row_content))

        return "\n".join(table_content)

    def _extract_table_content(self, table) -> str:
        """提取表格内容"""
        return self._extract_table_text(table, 0)

    def _is_heading(self, paragraph: str) -> bool:
        """判断段落是否是标题"""
        # 简单的启发式规则
        if len(paragraph) < 50:  # 标题通常较短
            # 检查常见标题模式
            heading_indicators = [
                '第', '章', '节', '条', '款', '项',  # 中文法律文档
                '一、', '二、', '三、', '四、', '五、',  # 中文序号
                '1.', '2.', '3.', '4.', '5.',  # 数字序号
                'I.', 'II.', 'III.', 'IV.', 'V.',  # 罗马数字
                'A.', 'B.', 'C.', 'D.', 'E.',  # 字母序号
            ]

            for indicator in heading_indicators:
                if paragraph.startswith(indicator):
                    return True

            # 检查是否全部是大写字母（英文标题）
            if paragraph.isupper() and len(paragraph) > 3:
                return True

        return False

    async def validate_file(self, file_path: str) -> bool:
        """验证Word文档完整性"""
        # 先进行基本验证
        if not await super().validate_file(file_path):
            return False

        try:
            # 尝试打开Word文档验证完整性
            doc = DocxDocument(file_path)

            # 检查是否有内容
            if len(doc.paragraphs) == 0 and len(doc.tables) == 0:
                return False

            # 尝试读取第一段
            if doc.paragraphs:
                _ = doc.paragraphs[0].text  # 尝试读取文本

            return True

        except Exception:
            return False