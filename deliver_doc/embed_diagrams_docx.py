from pathlib import Path
from docx import Document
from docx.shared import Inches
from docx.oxml import OxmlElement

OUT = Path("deliver_doc/output")
mapping = {
    "AIA-智能法律助手-SRS-v1.0-20260716.docx": ["srs_e2e.png"],
    "AIA-智能法律助手-HLD-v1.0-20260716.docx": [
        "hld_architecture.png", "hld_sequence.png", "hld_deployment.png"
    ],
}

def remove(paragraph):
    paragraph._element.getparent().remove(paragraph._element)

for name, images in mapping.items():
    source = OUT / name
    doc = Document(source)
    image_iter = iter(images)
    index = 0
    while index < len(doc.paragraphs):
        p = doc.paragraphs[index]
        if p.text.startswith(("flowchart ", "sequenceDiagram")):
            image = next(image_iter)
            # Mermaid code is split across individual Word paragraphs; remove it all.
            end = index
            while end < len(doc.paragraphs) and doc.paragraphs[end].text.strip() != "```":
                end += 1
            for old in list(doc.paragraphs[index:end + 1]):
                remove(old)
            anchor = doc.paragraphs[index] if index < len(doc.paragraphs) else doc.add_paragraph()
            run = anchor.insert_paragraph_before().add_run()
            run.add_picture(str(OUT / "images" / image), width=Inches(6.6))
            cap = anchor.insert_paragraph_before()
            cap.alignment = 1
            cap.add_run("可视化架构/流程图").bold = True
        else:
            index += 1
    target = source.with_name(source.stem + "-可视化.docx")
    doc.save(target)
    print(target)
